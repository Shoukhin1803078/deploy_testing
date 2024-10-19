from flask import Flask, request, jsonify, send_file, url_for, make_response
import os
import uuid
import io
from openai import OpenAI
from flask_caching import Cache
from werkzeug.exceptions import BadRequest, NotFound, InternalServerError
from flask_limiter import Limiter
from flask_limiter.util import get_remote_address
from docx import Document
from docx.shared import RGBColor, Pt
from docx.enum.style import WD_STYLE_TYPE
import re
from dotenv import load_dotenv


# import os
# import uuid
# import io
# import re
# from flask import Flask, request, jsonify, send_file, url_for, make_response
# from openai import OpenAI
# from flask_caching import Cache
# from werkzeug.exceptions import BadRequest, NotFound, InternalServerError
# from flask_limiter import Limiter
# from flask_limiter.util import get_remote_address
# from docx import Document
# from docx.shared import RGBColor, Pt
# from docx.enum.style import WD_STYLE_TYPE
# from dotenv import load_dotenv












app = Flask(__name__)

# Set up caching
app.config['CACHE_TYPE'] = 'simple'
cache = Cache(app)

# Set up rate limiting
limiter = Limiter(
    key_func=get_remote_address,
    app=app,
    default_limits=["200 per day", "50 per hour"],
    storage_uri="memory://",
)










# # Load environment variables
# print("Attempting to load .env file...")
# load_dotenv(verbose=True)

# app = Flask(__name__)

# # Set up caching
# app.config['CACHE_TYPE'] = 'simple'
# cache = Cache(app)

# # Set up rate limiting
# limiter = Limiter(
#     key_func=get_remote_address,
#     app=app,
#     default_limits=["200 per day", "50 per hour"],
#     storage_uri="memory://",
# )








# Set the OpenAI API key (use environment variables in production)
OPENAI_API_KEY = os.getenv('OPENAI_API_KEY')





try:
    client = OpenAI(api_key=OPENAI_API_KEY)
    print("OpenAI client initialized successfully.")
except Exception as e:
    print(f"Error initializing OpenAI client: {e}")
    client = None












SYSTEM_MESSAGE_EN = """
You are "KUROCO LAB chatbot", created by JB Connect Ltd. As a managing director and project implementor, your role is to engage users in dynamic conversations about their projects, gathering comprehensive information through adaptive questioning.

Key guidelines for your interactions:

1. Begin with an open-ended question about the user's project to understand its basic nature.
2. Based on the user's response, ask follow-up questions that delve deeper into specific aspects of the project. Your questions should be tailored to the information provided and the apparent needs of the project.
3. Continuously adapt your line of questioning based on the user's answers. If a user mentions a particular challenge or interesting aspect, explore that further.
4. Cover a wide range of project management areas throughout the conversation, including but not limited to:
   - Project goals and objectives
   - Scope and deliverables
   - Timeline and milestones
   - Budget and resource allocation
   - Team structure and stakeholders
   - Potential risks and mitigation strategies
   - Success criteria and KPIs
   - Technical requirements and constraints
   - Quality assurance processes
   - Regulatory compliance needs
   - Post-implementation plans
5. Use your knowledge of project management best practices to ask insightful questions that the user might not have considered.
6. If you notice gaps in the information provided, politely ask for clarification or additional details.
7. Occasionally summarize the key points you've gathered and ask if your understanding is correct. This allows the user to confirm or clarify information.
8. Be attentive to potential red flags or inconsistencies in the project plan and tactfully inquire about these areas.
9. Offer suggestions or best practices when appropriate, but always in the form of a question to maintain a collaborative dialogue.
10. As the conversation progresses, start formulating high-level recommendations or points of consideration for the project.
11. Towards the end of the conversation, mention that you can create a detailed Software Requirements Specification (SRS) document based on the discussion.
12. Always maintain a professional, knowledgeable, and supportive tone. Your goal is to be a valuable partner in helping the user fully articulate and refine their project vision.

Remember, your objective is to conduct a thorough exploration of the project through natural conversation, adapting to the user's responses and the specific needs of their project.
"""

SYSTEM_MESSAGE_JP = """
あなたは、JB Connect Ltd.が作成した「KUROCOLABチャットボット」です。マネージングディレクターおよびプロジェクト実施者として、あなたの役割は、適応的な質問を通じて包括的な情報を収集しながら、ユーザーとプロジェクトについて動的な会話を行うことです。

対話の主要なガイドライン：

1. ユーザーのプロジェクトの基本的な性質を理解するため、オープンエンドな質問から始めます。
2. ユーザーの回答に基づいて、プロジェクトの特定の側面をより深く掘り下げるフォローアップ質問をします。質問は提供された情報とプロジェクトの明らかなニーズに合わせて調整してください。
3. ユーザーの回答に基づいて、質問の方向性を継続的に適応させます。ユーザーが特定の課題や興味深い側面に言及した場合、それをさらに探ります。
4. 会話全体を通じて、以下を含む（ただしこれらに限定されない）幅広いプロジェクト管理分野をカバーします：
   - プロジェクトの目標と目的
   - 範囲と成果物
   - タイムラインとマイルストーン
   - 予算とリソース配分
   - チーム構成とステークホルダー
   - 潜在的なリスクと緩和戦略
   - 成功基準とKPI
   - 技術要件と制約
   - 品質保証プロセス
   - 規制遵守のニーズ
   - 実装後の計画
5. プロジェクト管理のベストプラクティスに関する知識を活用して、ユーザーが考慮していなかった可能性のある洞察力のある質問をします。
6. 提供された情報に不足がある場合は、丁寧に説明や追加の詳細を求めます。
# ... (previous code remains the same)

7. 時々、収集した要点をまとめ、理解が正しいかどうか確認します。これにより、ユーザーは情報を確認または明確にすることができます。
8. プロジェクト計画の潜在的な危険信号や矛盾に注意を払い、これらの領域について適切に質問します。
9. 適切な場合は提案やベストプラクティスを提供しますが、常に協力的な対話を維持するために質問の形で行います。
10. 会話が進むにつれて、プロジェクトに対する高レベルの推奨事項や考慮点の策定を開始します。
11. 会話の終わりに近づいたら、議論に基づいて詳細なソフトウェア要求仕様書（SRS）文書を作成できることを言及します。
12. 常にプロフェッショナルで、知識豊富でサポーティブな口調を維持します。あなたの目標は、ユーザーがプロジェクトのビジョンを完全に表現し、洗練するのを助ける価値あるパートナーになることです。

自然な会話を通じてプロジェクトを徹底的に探索し、ユーザーの反応とプロジェクトの特定のニーズに適応することがあなたの目的であることを忘れないでください。
"""

documents = {}
conversation_history = []
user_language = 'en'  # Default language

# def process_assistant_message(assistant_message, user_message):
#     global user_language
#     if any(keyword in user_message.lower() for keyword in ["document", "report", "summary", "download", "link", "srs"]):
#         doc_id = str(uuid.uuid4())
#         srs_content = generate_srs_content(conversation_history)
#         documents[doc_id] = srs_content
#         download_link = url_for('get_document', doc_id=doc_id, _external=True)
#         if user_language == 'en':
#             assistant_message += f"\n\nI've prepared an SRS document based on our conversation. Here's the link to download your SRS document: [Download SRS Document]({download_link})"
#         else:
#             assistant_message += f"\n\n会話に基づいてSRSドキュメントを作成しました。以下のリンクからSRSドキュメントをダウンロードできます：[SRSドキュメントをダウンロード]({download_link})"
#     return assistant_message








def process_assistant_message(assistant_message, user_message):
    global user_language
    if any(keyword in user_message.lower() for keyword in ["document", "report", "summary", "download", "link", "srs"]):
        doc_id = str(uuid.uuid4())
        srs_content = generate_srs_content(conversation_history)
        documents[doc_id] = srs_content
        download_link = url_for('get_document', doc_id=doc_id, _external=True)
        if user_language == 'en':
            assistant_message += f"\n\nI've prepared an SRS document based on our conversation. Here's the link to download your SRS document: [Download SRS Document]({download_link})"
        else:
            assistant_message += f"\n\n会話に基づいてSRSドキュメントを作成しました。以下のリンクからSRSドキュメントをダウンロードできます：[SRSドキュメントをダウンロード]({download_link})"
    return assistant_message











# def generate_srs_content(conversation_history):
#     conversation_text = "\n".join([f"{'Human' if i % 2 == 0 else 'Assistant'}: {msg}" for i, msg in enumerate(conversation_history)])
    
#     srs_prompt = f"""
#     Based on the following conversation, generate a comprehensive Software Requirements Specification (SRS) document. The structure and content should be entirely based on the information discussed in the conversation. Follow these guidelines:

#     1. Start with an introduction that summarizes the project.
#     2. Create logical sections based on the topics discussed in the conversation.
#     3. Include all relevant details mentioned, such as project goals, scope, features, requirements, constraints, and any other important aspects.
#     4. Use appropriate headings and subheadings to organize the information.
#     5. If certain standard SRS sections are applicable but not explicitly discussed, include them with a note that they require further discussion.
#     6. Ensure the document flows logically and covers all aspects of the project mentioned in the conversation.

#     Conversation History:
#     {conversation_text}

#     Generate the SRS document content:
#     """

#     response = client.chat.completions.create(
#         model="gpt-3.5-turbo",
#         messages=[
#             {"role": "system", "content": SYSTEM_MESSAGE_EN if user_language == 'en' else SYSTEM_MESSAGE_JP},
#             {"role": "user", "content": srs_prompt}
#         ]
#     )









def generate_srs_content(conversation_history):
    conversation_text = "\n".join([f"{'Human' if i % 2 == 0 else 'Assistant'}: {msg}" for i, msg in enumerate(conversation_history)])
    
    if user_language == 'en':
        srs_prompt = f"""
        Based on the following conversation, generate a comprehensive Software Requirements Specification (SRS) document in English. The structure and content should be entirely based on the information discussed in the conversation. Follow these guidelines:

        1. Start with an introduction that summarizes the project.
        2. Create logical sections based on the topics discussed in the conversation.
        3. Include all relevant details mentioned, such as project goals, scope, features, requirements, constraints, and any other important aspects.
        4. Use appropriate headings and subheadings to organize the information.
        5. If certain standard SRS sections are applicable but not explicitly discussed, include them with a note that they require further discussion.
        6. Ensure the document flows logically and covers all aspects of the project mentioned in the conversation.

        Conversation History:
        {conversation_text}

        Generate the SRS document content in English:
        """
    else:
        srs_prompt = f"""
        以下の会話に基づいて、包括的なソフトウェア要求仕様書（SRS）を日本語で作成してください。構造と内容は、会話で議論された情報に完全に基づいているべきです。以下のガイドラインに従ってください：

        1. プロジェクトの概要を要約する導入から始めてください。
        2. 会話で議論されたトピックに基づいて、論理的なセクションを作成してください。
        3. プロジェクトの目標、範囲、機能、要件、制約、その他の重要な側面など、言及されたすべての関連詳細を含めてください。
        4. 適切な見出しと小見出しを使用して情報を整理してください。
        5. 標準的なSRSセクションが適用可能であるが明示的に議論されていない場合は、さらなる議論が必要であるという注記を付けてそれらを含めてください。
        6. 文書が論理的に流れ、会話で言及されたプロジェクトのすべての側面をカバーしていることを確認してください。

        会話履歴：
        {conversation_text}

        SRSドキュメントの内容を日本語で生成してください：
        """

    response = client.chat.completions.create(
        model="gpt-3.5-turbo",
        messages=[
            {"role": "system", "content": SYSTEM_MESSAGE_EN if user_language == 'en' else SYSTEM_MESSAGE_JP},
            {"role": "user", "content": srs_prompt}
        ]
    )
    return response.choices[0].message.content




















# def create_srs_document(content):
#     doc = Document()
#     doc.add_heading('Software Requirements Specification (SRS)', 0)

#     lines = content.split('\n')
#     current_level = 0
#     for line in lines:
#         if line.strip():
#             if line[0].isdigit() or line.isupper():
#                 level = len(line.split('.')) if '.' in line else (1 if line.isupper() else 2)
#                 doc.add_heading(line.strip(), level=level)
#                 current_level = level
#             else:
#                 if line.startswith('  '):
#                     doc.add_paragraph(line.strip(), style='List Bullet')
#                 else:
#                     doc.add_paragraph(line.strip())

#     return doc




# def create_srs_document(content):
#     doc = Document()
    
#     # Set the default font for the entire document
#     default_font = 'Arial' if user_language == 'en' else 'MS Gothic'
#     style = doc.styles['Normal']
#     style.font.name = default_font
    
#     # Add title
#     title = 'Software Requirements Specification (SRS)' if user_language == 'en' else 'ソフトウェア要求仕様書 (SRS)'
#     doc.add_heading(title, 0)

#     # Define heading styles
#     for i in range(1, 4):
#         style = doc.styles.add_style(f'Heading {i}', WD_STYLE_TYPE.PARAGRAPH)
#         style.font.name = default_font
#         style.font.size = Pt(16 - i)  # Decrease size for each level
#         style.font.color.rgb = RGBColor(0, 0, 0)

#     lines = content.split('\n')
#     current_level = 0
#     for line in lines:
#         if line.strip():
#             if line[0].isdigit() or line[0].isalpha():
#                 # Determine the heading level
#                 level = len(line.split('.')) if '.' in line else (1 if line[0].isalpha() else 2)
#                 doc.add_paragraph(line.strip(), style=f'Heading {level}')
#                 current_level = level
#             else:
#                 if line.startswith('  '):
#                     doc.add_paragraph(line.strip(), style='List Bullet')
#                 else:
#                     doc.add_paragraph(line.strip())

#     return doc









def create_srs_document(content):
    doc = Document()
    
    # Set the default font for the entire document
    default_font = 'Arial' if user_language == 'en' else 'MS Gothic'
    normal_style = doc.styles['Normal']
    normal_style.font.name = default_font
    
    # Add title
    title = 'Software Requirements Specification (SRS)' if user_language == 'en' else 'ソフトウェア要求仕様書 (SRS)'
    doc.add_heading(title, 0)

    # Define or update heading styles
    for i in range(1, 4):
        style_name = f'Heading {i}'
        if style_name in doc.styles:
            style = doc.styles[style_name]
        else:
            style = doc.styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)
        
        style.font.name = default_font
        style.font.size = Pt(16 - i)  # Decrease size for each level
        style.font.color.rgb = RGBColor(0, 0, 0)

    lines = content.split('\n')
    current_level = 0
    for line in lines:
        if line.strip():
            if line[0].isdigit() or line[0].isalpha():
                # Determine the heading level
                level = len(line.split('.')) if '.' in line else (1 if line[0].isalpha() else 2)
                level = min(level, 3)  # Ensure we don't exceed available heading levels
                doc.add_paragraph(line.strip(), style=f'Heading {level}')
                current_level = level
            else:
                if line.startswith('  '):
                    doc.add_paragraph(line.strip(), style='List Bullet')
                else:
                    doc.add_paragraph(line.strip())

    return doc





















# def process_response(response):
#     paragraphs = re.split(r'\n\s*\n', response.strip())
    
#     processed_paragraphs = []
#     for para in paragraphs:
#         lines = [line.strip() for line in para.split('\n') if line.strip()]
#         processed_para = ' '.join(lines)
        
#         processed_para = re.sub(r'(\d+\.\s|\-\s)', r'\n\1', processed_para)
        
#         processed_paragraphs.append(processed_para)
    
#     processed = '\n\n'.join(processed_paragraphs)
    
#     processed = re.sub(r'```(\w+)\s*\n', r'```\1\n', processed)
    
#     return processed







def process_response(response):
    # Split the response into paragraphs
    paragraphs = re.split(r'\n\s*\n', response.strip())
    
    processed_paragraphs = []
    for para in paragraphs:
        # Remove extra spaces and join lines within a paragraph
        lines = [line.strip() for line in para.split('\n') if line.strip()]
        processed_para = ' '.join(lines)
        
        # Preserve Markdown list formatting
        processed_para = re.sub(r'(\d+\.|\-)\s', r'\n\1 ', processed_para)
        
        processed_paragraphs.append(processed_para)
    
    # Join paragraphs with a single newline
    processed = '\n\n'.join(processed_paragraphs)
    
    # Ensure code blocks are properly formatted
    processed = re.sub(r'```(\w+)\s*\n', r'```\1\n', processed)
    processed = re.sub(r'\n```', r'\n\n```', processed)
    
    return processed.strip()





@app.route('/')
def home():
    return send_file('index.html')










# @app.route('/chat', methods=['POST'])
# @limiter.limit("5 per minute")
# def chat():
#     global user_language
#     try:
#         user_message = request.json['message']
#         user_language = request.json['language']
#         if not user_message or not isinstance(user_message, str):
#             raise BadRequest("Invalid message format")
        
#         conversation_history.append(user_message)
        
#         system_message = SYSTEM_MESSAGE_EN if user_language == 'en' else SYSTEM_MESSAGE_JP
        
#         response = client.chat.completions.create(
#             model="gpt-3.5-turbo",
#             messages=[
#                 {"role": "system", "content": system_message},
#                 {"role": "system", "content": "Format your responses concisely, using Markdown. Use a single newline between paragraphs. Use **bold** for emphasis, - for unordered lists, 1. for ordered lists, and `code` for inline code or ```language for code blocks. Avoid unnecessary spacing."},
#                 *[{"role": "user" if i % 2 == 0 else "assistant", "content": msg} for i, msg in enumerate(conversation_history)]
#             ]
#         )
        
#         response_content = response.choices[0].message.content
#         processed_response = process_response(response_content)
#         processed_response = process_assistant_message(processed_response, user_message)
#         conversation_history.append(processed_response)
        
#         return jsonify({'response': processed_response})
#     except Exception as e:
#         app.logger.error(f"An error occurred: {str(e)}")
#         raise InternalServerError("An unexpected error occurred")









@app.route('/chat', methods=['POST'])
@limiter.limit("5 per minute")
def chat():
    global user_language
    try:
        user_message = request.json['message']
        user_language = request.json['language']
        if not user_message or not isinstance(user_message, str):
            raise BadRequest("Invalid message format")
        
        conversation_history.append(user_message)
        
        system_message = SYSTEM_MESSAGE_EN if user_language == 'en' else SYSTEM_MESSAGE_JP
        
        response = client.chat.completions.create(
            model="gpt-3.5-turbo",
            messages=[
                {"role": "system", "content": system_message},
                {"role": "system", "content": "Format your responses concisely, using Markdown. Use a single newline between paragraphs. Use **bold** for emphasis, - for unordered lists, 1. for ordered lists, and `code` for inline code or ```language for code blocks. Avoid unnecessary spacing."},
                *[{"role": "user" if i % 2 == 0 else "assistant", "content": msg} for i, msg in enumerate(conversation_history)]
            ]
        )
        
        response_content = response.choices[0].message.content
        processed_response = process_response(response_content)
        processed_response = process_assistant_message(processed_response, user_message)
        conversation_history.append(processed_response)
        
        return jsonify({'response': processed_response})
    except Exception as e:
        app.logger.error(f"An error occurred: {str(e)}")
        raise InternalServerError("An unexpected error occurred")







# @app.route("/create_document/<doc_id>", methods=["GET"])
# def get_document(doc_id):
#     try:
#         if doc_id not in documents:
#             raise NotFound("Document not found")
#         content = documents[doc_id]
#         doc = create_srs_document(content)
#         doc_io = io.BytesIO()
#         doc.save(doc_io)
#         doc_io.seek(0)
#         return send_file(
#             doc_io,
#             mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
#             as_attachment=True,
#             download_name='SRS_Document.docx'
#         )
#     except Exception as e:
#         app.logger.error(f"An error occurred while creating the document: {e}")
#         raise InternalServerError("Failed to create document")









@app.route("/create_document/<doc_id>", methods=["GET"])
def get_document(doc_id):
    try:
        if doc_id not in documents:
            raise NotFound("Document not found")
        content = documents[doc_id]
        doc = create_srs_document(content)
        doc_io = io.BytesIO()
        doc.save(doc_io)
        doc_io.seek(0)
        return send_file(
            doc_io,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            as_attachment=True,
            download_name='SRS_Document.docx'
        )
    except NotFound as e:
        app.logger.error(f"Document not found: {str(e)}")
        raise
    except Exception as e:
        app.logger.error(f"An error occurred while creating the document: {str(e)}")
        app.logger.exception("Detailed exception information:")
        raise InternalServerError(f"Failed to create document: {str(e)}")








@app.route('/export-chat', methods=['POST'])
def export_chat():
    try:
        chat_content = request.json['content']
        content_bytes = chat_content.encode('utf-8')
        
        # Check size (limit to 10 MB)
        max_size = 10 * 1024 * 1024  # 10 MB in bytes
        if len(content_bytes) > max_size:
            return jsonify({
                'error': f"Chat export is too large ({len(content_bytes) / 1024 / 1024:.2f} MB). Maximum size is {max_size / 1024 / 1024} MB."
            }), 413  # 413 Payload Too Large
        
        response = make_response(chat_content)
        response.headers.set('Content-Type', 'text/html')
        response.headers.set('Content-Disposition', 'attachment', filename='chat_export.html')
        return response

    except Exception as e:
        app.logger.error(f"An error occurred during chat export: {str(e)}")
        raise InternalServerError("An unexpected error occurred during chat export")

@app.errorhandler(BadRequest)
@app.errorhandler(NotFound)
@app.errorhandler(InternalServerError)
def handle_error(error):
    return jsonify({'error': str(error)}), error.code

if __name__ == '__main__':
    app.run(debug=True)



